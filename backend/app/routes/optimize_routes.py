import os
import docx
from io import BytesIO
from flask import Blueprint, request, jsonify, send_file, current_app
# We will use your find_and_replace_in_doc function
from app.utils.file_utils import find_and_replace_in_doc, extract_json_from_response
from app.utils.ai_helpers import client
import json # <-- Make sure you have this import for the new route
import requests


OLLAMA_BASE_URL= os.getenv("OLLAMA_BASE_URL")

optimize_bp = Blueprint('optimize', __name__)

# --- NO CHANGES TO THIS SECTION ---
def insert_paragraph_after(paragraph, text, style):
    """A helper function to insert a new paragraph after a given one."""
    new_p = docx.text.paragraph.Paragraph(paragraph._p.getparent()._new_p(), paragraph._parent)
    if text:
        new_p.add_run(text)
    if style:
        new_p.style = style
    paragraph._p.addnext(new_p._p)
    return new_p

@optimize_bp.route("/optimize_resume", methods=["POST"])
def optimize_resume_route():
    """Optimize resume based on user answers to questions"""
    data = request.get_json()
    jd_text = data.get('jd_text')
    answers = data.get('answers')
    file_id = data.get('file_id')
    old_score = data.get('old_score') 
    
    if not all([jd_text, answers, file_id, old_score is not None]):
        return jsonify({"error": "JD, answers, file_id, and old_score are required."}), 400

    target_score = round(old_score * 1.10)

    file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], file_id)
  
    if not os.path.exists(file_path):
        return jsonify({"error": "Original resume file not found."}), 404
        
    if not file_path.lower().endswith('.docx'):
        return jsonify({"error": "Optimization is for .docx only."}), 400
    
    original_doc = docx.Document(file_path)
    original_resume_text = "\n".join([p.text for p in original_doc.paragraphs if p.text.strip()])
    
    formatted_answers = []
    for q, a in answers.items():
        if a.strip():
            formatted_answers.append(f"Question: {q}\nAnswer: {a}\n")
    
    # Define the exact JSON structure we expect
    json_structure = {
        "skills_to_add": ["Skill1", "Skill2", "Skill3"],
        "project_enhancements": [
            {
                "anchor_bullet": "Existing bullet point text",
                "new_bullet_to_add": "• New enhanced bullet point"
            }
        ],
        "bullet_point_changes": [
            {
                "find": "Original text to find",
                "replace": "Enhanced replacement text"
            }
        ]
    }
    
    prompt = f"""
CRITICAL INSTRUCTION: YOU MUST RETURN ONLY VALID JSON. DO NOT INCLUDE ANY OTHER TEXT, ANALYSIS, OR EXPLANATIONS.

Analyze this resume against the job description and return ONLY a JSON object with this exact structure:

{json.dumps(json_structure, indent=2)}

JOB DESCRIPTION:
{jd_text}

CANDIDATE ANSWERS:
{"".join(formatted_answers)}

RESUME CONTENT:
{original_resume_text[:12000]}

ANALYSIS RULES:
1. Tailoring Score (35% weight): Match JD requirements against resume content
2. Content Score (20% weight): Evaluate relevance and evidence for JD requirements  
3. Format Score (15% weight): Check ATS compatibility and professional layout
4. Sections Score (15% weight): Verify essential sections are present
5. Style Score (15% weight): Assess language quality and professionalism

RETURN ONLY JSON. NO OTHER TEXT.
"""
    
    try:
        payload = {
            "model_name": "mistral:7b-instruct",  # Try a different model if mistral isn't working
            "prompt": prompt,
            "actual_data": original_resume_text[:12000],
            "temperature": 0.1,  # Lower temperature for more deterministic output
            "top_p": 0.9,
            "max_tokens": 2000,
            "top_k": 40,
            "repeat_penalty": 1.1,
            "num_ctx": 4096,
            "num_thread": 4,
            "stop": ["\n\n", "Human:", "```", "Analysis:", "Here is", "The resume", "Name:"],
            "conversation_history": []
        }
        
        current_app.logger.info(f"Sending request to Ollama with prompt: {prompt[:200]}...")
        
        response = requests.post(OLLAMA_BASE_URL, json=payload)

        if response.status_code != 200:
            return jsonify({"error": f"Ollama API error: {response.text}"}), 500
        
        result = response.json()
        
        # Check if the response contains the expected structure
        if not result.get("success", False):
            return jsonify({"error": f"Ollama API returned unsuccessful response: {result.get('message', 'Unknown error')}"}), 500
            
        response_text = result.get("response", "")

        current_app.logger.info(f"Raw AI response: {response_text[:500]}...")
        
        # Enhanced JSON extraction with better error handling
        cleaned_text = response_text.strip()
        
        # Remove markdown code blocks if present
        if cleaned_text.startswith('```json'):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.startswith('```'):
            cleaned_text = cleaned_text[3:]
        if cleaned_text.endswith('```'):
            cleaned_text = cleaned_text[:-3]
        cleaned_text = cleaned_text.strip()
        
        # Try to parse the JSON
        try:
            optimization_data = json.loads(cleaned_text)
        except json.JSONDecodeError as e:
            current_app.logger.error(f"JSON parsing failed. Cleaned response: {cleaned_text[:500]}")
            
            # Try to extract JSON using more aggressive methods
            json_match = re.search(r'\{.*\}', cleaned_text, re.DOTALL)
            if json_match:
                try:
                    optimization_data = json.loads(json_match.group())
                    current_app.logger.info("Successfully extracted JSON using regex fallback")
                except json.JSONDecodeError:
                    # If it's still not JSON, try a different approach
                    return jsonify({
                        "error": "AI returned analysis instead of JSON", 
                        "raw_preview": cleaned_text[:200] + "...",
                        "solution": "Try using a different model or simplifying the prompt"
                    }), 500
            else:
                return jsonify({
                    "error": "AI returned analysis instead of JSON", 
                    "raw_preview": cleaned_text[:200] + "...",
                    "solution": "Try using a different model or simplifying the prompt"
                }), 500

        # Validate the structure
        required_keys = ["skills_to_add", "project_enhancements", "bullet_point_changes"]
        if not all(key in optimization_data for key in required_keys):
            return jsonify({"error": "AI returned incomplete analysis structure"}), 500

        skills_to_add = optimization_data.get("skills_to_add", [])
        project_enhancements = optimization_data.get("project_enhancements", [])
        bullet_changes = optimization_data.get("bullet_point_changes", []) 
        
        if not skills_to_add and not project_enhancements and not bullet_changes:
            return jsonify({"error": "The AI could not generate any optimizations."}), 400

        modified_doc = docx.Document(file_path)
        
        if skills_to_add:
            for p in modified_doc.paragraphs:
                if p.text.strip().upper().startswith("SKILLS"):
                    if p.text.strip() and not p.text.strip().endswith(('|', '.', ',')):
                         p.text += ' |'
                    p.text += " " + " | ".join(skills_to_add)
                    break

        if project_enhancements:
            for enhancement in project_enhancements:
                anchor_text = enhancement.get("anchor_bullet")
                new_bullet_text = enhancement.get("new_bullet_to_add")

                if not anchor_text or not new_bullet_text:
                    continue

                for i, p in enumerate(modified_doc.paragraphs):
                    if anchor_text.strip() in p.text.strip():
                        insert_paragraph_after(p, new_bullet_text, p.style)
                        break

        if bullet_changes:
            changes = bullet_changes 
            modified_doc = find_and_replace_in_doc(modified_doc, changes)
        
        doc_io = BytesIO()
        modified_doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            as_attachment=True,
            download_name='Optimized_Resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        current_app.logger.error(f"Optimization failed: {str(e)}")
        current_app.logger.error(f"AI response was: {locals().get('response_text', 'Not available')}")
        return jsonify({"error": f"Optimization failed: {str(e)}"}), 500
# ========================================================================
# === THIS IS THE NEW, CORRECTED ROUTE FOR THE "DESIGN IN BUILDER" BUTTON ===
# ========================================================================

@optimize_bp.route("/parse_final_resume_to_json", methods=["POST"])
def parse_final_resume_to_json():
    """
    Receives the FINAL, already-optimized .docx file from the frontend.
    Its only job is to extract the text and parse it into JSON.
    This guarantees the text is identical to the downloaded file.
    """
    if 'optimized_resume_file' not in request.files:
        return jsonify({"error": "Optimized resume file is required."}), 400

    file = request.files['optimized_resume_file']

    try:
        # Re-import utility here to be safe
        from app.utils.file_utils import extract_text_builder
        # This gets the text from the FINAL .docx file sent from the frontend
        final_optimized_text = extract_text_builder(file, file.filename)
    except Exception as e:
        return jsonify({"error": f"Failed to extract text from optimized file: {str(e)}"}), 500
    
    # The JSON structure your TemplateBuilder expects
    target_json_structure = {
        "Name": "Full Name",
        "jobTitle": "Job Title",
        "email": "email@example.com",
        "phone": "123-456-7890",
        "location": "City, Country",
        "linkedin": "https://linkedin.com/...",
        "github": "https://github.com/...",
        "summary": "Professional summary...",
        "objective": "Career Objectives",
        "experience": [
            {
                "jobTitle": "Job Title",
                "company": "Company Name",
                "startDate": "Month Year",
                "endDate": "Month Year",
                "description": ["Point 1", "Point 2"]
            }
        ],
        "internship": [
            {
                "role": "Job Title",
                "company": "Company Name",
                "startDate": "Month Year",
                "endDate": "Month Year",
                "description": ["Point 1", "Point 2"]
            }
        ],
        "education": [
            {
                "degree": "Degree Name",
                "school": "School Name",
                "level": "type of education level",
                "startDate": "Year",
                "endDate": "Year",
                "cgpa": "X.XX/4.0"  
            }
        ],
        "skills": ["Skill1", "Skill2"],
        "languages": ["Language1", "Language2"],
        "interests": ["Interest1", "Interest2"],
        "achievements": [
            {
                "title": "Achievement Title",
                "description": ["Point 1", "Point 2"]
            }
        ],
        "projects": [
            {
                "title": "Project Title",
                "startDate": "Month Year",
                "endDate": "Month Year",
                "tech": "Tools/Tech Stack",
                "description": ["Point 1", "Point 2"]
            }
        ],
        "certifications": [
            {
                "name": "Certification name",
                "issuer": "name of the issuer",
                "Date": "Month Year"
            }
        ]
    }

    final_prompt = f"""
CRITICAL INSTRUCTION: YOU MUST RETURN ONLY VALID JSON. DO NOT INCLUDE ANY OTHER TEXT, ANALYSIS, OR EXPLANATIONS.

You are an expert data extractor. Your task is to parse the following complete resume text and structure it into a clean JSON object.
Do not invent any information. Extract the data exactly as it appears in the text.

RESUME TEXT TO PARSE:
---
{final_optimized_text[:15000]}
---

REQUIREMENTS:
1. Extract data exactly as it appears in the resume text
2. Do not invent any information - use only what's present
3. Return ONLY valid JSON, no explanatory text
4. Follow this exact structure:
{json.dumps(target_json_structure, indent=2)}

Now, generate the JSON object based on the provided text.
"""

    try:
        payload = {
            "model_name": "mistral:7b-instruct",
            "prompt": final_prompt,
            "actual_data": final_optimized_text[:15000],
            "temperature": 0.1,
            "top_p": 0.9,
            "max_tokens": 2500,
            "top_k": 40,
            "repeat_penalty": 1.1,
            "num_ctx": 4096,
            "num_thread": 4,
            "stop": ["\n\n", "Human:", "```", "Analysis:", "Here is", "The resume"],
            "conversation_history": []
        }
        
        current_app.logger.info(f"Sending parse request to Ollama with text length: {len(final_optimized_text)}")
        
        response = requests.post(OLLAMA_BASE_URL, json=payload)  # 2 minute timeout

        if response.status_code != 200:
            current_app.logger.error(f"Ollama API error: {response.status_code} - {response.text}")
            return jsonify({"error": f"Ollama API error: {response.text}"}), 500
    
        result = response.json()
        
        # Check if the response contains the expected structure
        if not result.get("success", False):
            return jsonify({"error": f"Ollama API returned unsuccessful response: {result.get('message', 'Unknown error')}"}), 500
            
        response_text = result.get("response", "")

        current_app.logger.info(f"Raw AI response: {response_text[:500]}...")
        
        # Enhanced JSON extraction with better error handling
        cleaned_text = response_text.strip()
        
        # Remove markdown code blocks if present
        if cleaned_text.startswith('```json'):
            cleaned_text = cleaned_text[7:]
        if cleaned_text.startswith('```'):
            cleaned_text = cleaned_text[3:]
        if cleaned_text.endswith('```'):
            cleaned_text = cleaned_text[:-3]
        cleaned_text = cleaned_text.strip()
        
        # Try to parse the JSON
        try:
            final_json_data = json.loads(cleaned_text)
        except json.JSONDecodeError as e:
            current_app.logger.error(f"JSON parsing failed. Cleaned response: {cleaned_text[:500]}")
            
            # Try to extract JSON using more aggressive methods
            json_match = re.search(r'\{.*\}', cleaned_text, re.DOTALL)
            if json_match:
                try:
                    final_json_data = json.loads(json_match.group())
                    current_app.logger.info("Successfully extracted JSON using regex fallback")
                except json.JSONDecodeError:
                    return jsonify({
                        "error": "AI returned invalid JSON format", 
                        "raw_preview": cleaned_text[:200] + "...",
                        "hint": "The AI might be analyzing the resume content instead of returning JSON."
                    }), 500
            else:
                return jsonify({
                    "error": "AI returned analysis instead of JSON format", 
                    "raw_preview": cleaned_text[:200] + "...",
                    "hint": "The prompt may need stronger instructions to return only JSON."
                }), 500
        
        return jsonify(final_json_data)
        
    except requests.exceptions.Timeout:
        return jsonify({"error": "Resume parsing timed out after 2 minutes"}), 408
    except Exception as e:
        current_app.logger.error(f"Final JSON parsing failed: {str(e)}")
        return jsonify({"error": f"Failed to generate resume JSON for builder: {str(e)}"}), 500


@optimize_bp.route("/generate-or-enhance-summary", methods=["POST"])
def generate_or_enhance_summary():
    """
    Generate or enhance a professional summary based on the user's input.
    - If 'existingSummary' is provided, enhances only that text.
    - If 'existingSummary' is empty, generates a new summary using 'experience', 'skills', and 'projects'.
    Expects a JSON payload with 'existingSummary', 'experience', 'skills', and 'projects' fields.
    Returns a JSON object with a 'summary' field containing the generated or enhanced summary.
    """
    data = request.get_json()
    existing_summary = data.get('existingSummary', '').strip()
    experience = data.get('experience', [])
    skills = data.get('skills', [])
    projects = data.get('projects', [])

    # Construct prompt based on whether existing_summary is provided
    if existing_summary:
        # Enhance mode: Only enhance the existing summary
        prompt = f"""CRITICAL INSTRUCTION: RETURN ONLY VALID JSON. DO NOT INCLUDE ANY OTHER TEXT.

ENHANCE PROFESSIONAL SUMMARY:

CONTEXT:
Existing Summary: {existing_summary}

REQUIREMENTS:
- Enhance the summary to be more concise, professional, and impactful
- Maintain the original message and tone
- Do not add new information beyond what is provided
- Keep it concise (3-5 sentences, max 150 words)
- Use professional tone suitable for a resume
- RETURN ONLY JSON: {{"summary": "enhanced summary text here"}}

EXAMPLE OUTPUT:
{{"summary": "Accomplished Software Engineer with over 5 years of experience, delivering high-quality web applications. Enhanced expertise in modern frameworks, driving system efficiency through scalable solutions."}}

Enhanced summary:"""
        
        # Use existing summary as actual_data
        actual_data = existing_summary
        
    else:
        # Generate mode: Create a new summary using experience, skills, and projects
        if not any([experience, skills, projects]):
            return jsonify({"error": "At least one of experience, skills, or projects is required when generating a new summary."}), 400

        # Format experience for the prompt
        formatted_experience = []
        for exp in experience:
            formatted_experience.append(
                f"Role: {exp.get('jobTitle', '')}\n"
                f"Company: {exp.get('company', '')}\n"
                f"Duration: {exp.get('startDate', '')} - {exp.get('endDate', 'Present')}\n"
                f"Responsibilities: {exp.get('description', '')}\n"
            )

        # Format skills
        formatted_skills = ", ".join(skills) if skills else "None provided"

        # Format projects
        formatted_projects = []
        for proj in projects:
            formatted_projects.append(
                f"Project: {proj.get('title', '')}\n"
                f"Duration: {proj.get('startDate', '')} - {proj.get('endDate', 'Present')}\n"
                f"Description: {proj.get('description', '')}\n"
                f"Technologies: {proj.get('tech', '')}\n"
            )

        prompt = f"""CRITICAL INSTRUCTION: RETURN ONLY VALID JSON. DO NOT INCLUDE ANY OTHER TEXT.

GENERATE PROFESSIONAL SUMMARY:

CONTEXT:
Experience: {"".join(formatted_experience) if formatted_experience else "None provided"}
Skills: {formatted_skills}
Projects: {"".join(formatted_projects) if formatted_projects else "None provided"}

REQUIREMENTS:
- Generate professional summary reflecting experience, skills, and projects
- Keep concise (3-4 sentences, max 60-70 words)
- Use professional tone suitable for resume
- Incorporate key achievements from experience and projects
- Highlight relevant skills
- RETURN ONLY JSON: {{"summary": "generated summary text here"}}

EXAMPLE OUTPUT:
{{"summary": "Results-driven Software Engineer with over 5 years of experience developing scalable web applications at leading tech firms. Proficient in Java, Python, and cloud technologies, with a proven track record of delivering high-impact projects that enhance system performance."}}

Generated summary:"""
        
        # Combine experience, skills, and projects for actual_data
        actual_data = f"Experience: {formatted_experience}\nSkills: {formatted_skills}\nProjects: {formatted_projects}"
    
    try:
        payload = {
            "model_name": "mistral:7b-instruct",
            "prompt": prompt,
            "actual_data": actual_data[:5000],  # Limit to 5000 characters
            "temperature": 0.3,
            "top_p": 0.9,
            "max_tokens": 300,
            "top_k": 40,
            "repeat_penalty": 1.1,
            "num_ctx": 2048,
            "num_thread": 4,
            "stop": ["\n\n", "Human:", "```", "Analysis:", "Here is", "The summary"],
            "conversation_history": []
        }
        
        current_app.logger.info(f"Sending summary request to Ollama with prompt: {prompt[:200]}...")
        
        response = requests.post(OLLAMA_BASE_URL, json=payload)  # 1.5 minute timeout
        
        if response.status_code != 200:
            current_app.logger.error(f"Ollama API error: {response.status_code} - {response.text}")
            return jsonify({"error": f"Ollama API error: {response.text}"}), 500
            
        result = response.json()
        
        # Check if the response contains the expected structure
        if not result.get("success", False):
            return jsonify({"error": f"Ollama API returned unsuccessful response: {result.get('message', 'Unknown error')}"}), 500
            
        response_text = result.get("response", "")
        
        current_app.logger.info(f"AI response: {response_text[:200]}...")
        
        # Extract JSON from response
        try:
            # Try to find JSON in the response
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            if json_start != -1 and json_end != -1:
                json_str = response_text[json_start:json_end]
                summary_data = json.loads(json_str)
            else:
                # If no JSON found, try to parse directly
                summary_data = json.loads(response_text)
        except json.JSONDecodeError as e:
            current_app.logger.error(f"JSON parsing failed. Response: {response_text}")
            return jsonify({"error": f"Failed to parse JSON from AI response: {str(e)}"}), 500
        
        if not summary_data.get("summary"):
            return jsonify({"error": "Failed to generate a valid summary."}), 500

        return jsonify(summary_data)
        
    except requests.exceptions.Timeout:
        return jsonify({"error": "Summary generation timed out after 1.5 minutes"}), 408
    except Exception as e:
        current_app.logger.error(f"Summary generation/enhancement failed: {str(e)}")
        return jsonify({"error": f"Failed to process summary: {str(e)}"}), 500
        