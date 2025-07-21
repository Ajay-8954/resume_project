# import fitz
# import docx
# from werkzeug.utils import secure_filename
# import json
# from thefuzz import fuzz

# def extract_text_builder(file, filename):
#     filename = filename.lower()
#     if filename.endswith(".pdf"):
#         doc = fitz.open(stream=file.read(), filetype="pdf")
#         return " ".join(page.get_text() for page in doc)
#     elif filename.endswith(".docx"):
#         return "\n".join(para.text for para in docx.Document(file).paragraphs)
#     elif filename.endswith(".txt"):
#         return file.read().decode("utf-8", errors="ignore")
#     return ""

# def extract_text(file_storage):
#     filename = secure_filename(file_storage.filename)
#     file_storage.seek(0)
#     if filename.lower().endswith(".pdf"):
#         return " ".join(page.get_text() for page in fitz.open(stream=file_storage.read(), filetype="pdf"))
#     elif filename.lower().endswith(".docx"):
#         return "\n".join(p.text for p in docx.Document(file_storage).paragraphs)
#     return ""

# def extract_json_from_response(response_text):
#     try:
#         return json.loads(response_text)
#     except json.JSONDecodeError:
#         if '```' in response_text:
#             json_str = response_text.split('```')[1].split('```')[0]
#             return json.loads(json_str)
#         else:
#             json_str = response_text[response_text.find('{'):response_text.rfind('}')+1]
#             return json.loads(json_str)

# def find_and_replace_in_doc(doc, changes):
#     for change in changes:
#         find_text = change.get('find', '').strip().lstrip('•*-').strip()
#         replace_text = change.get('replace', '').strip()
#         for p in doc.paragraphs:
#             if find_text in p.text:
#                 p.text = p.text.replace(find_text, replace_text)
#     return doc




import fitz
import docx
from werkzeug.utils import secure_filename
import json
from thefuzz import fuzz
import re # Added for a more robust JSON extraction

def extract_text_builder(file, filename):
    filename = filename.lower()
    if filename.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return " ".join(page.get_text() for page in doc)
    elif filename.endswith(".docx"):
        return "\n".join(para.text for para in docx.Document(file).paragraphs)
    elif filename.endswith(".txt"):
        return file.read().decode("utf-8", errors="ignore")
    return ""

def extract_text(file_storage):
    filename = secure_filename(file_storage.filename)
    file_storage.seek(0)
    if filename.lower().endswith(".pdf"):
        return " ".join(page.get_text() for page in fitz.open(stream=file_storage.read(), filetype="pdf"))
    elif filename.lower().endswith(".docx"):
        return "\n".join(p.text for p in docx.Document(file_storage).paragraphs)
    return ""

def extract_json_from_response(response_text):
    # Using a slightly more robust method to find JSON, but it's a minor improvement.
    try:
        return json.loads(response_text)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", response_text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                raise ValueError("Could not extract valid JSON from response.")
        raise ValueError("Could not find a JSON object in the response.")


def find_and_replace_in_doc(doc, changes):
    for change in changes:
        find_text = change.get('find', '').strip().lstrip('•*-').strip()
        replace_text = change.get('replace', '').strip()

        # Iterate through all paragraphs in the document
        for p in doc.paragraphs:
            # <<< --- START OF THE FIX --- >>>
            # This is the new "guard clause".
            # It checks for a hyperlink in the paragraph's XML. If found, it skips this paragraph.
            if p._p.xpath("./w:hyperlink"):
                continue
            # <<< ---  END OF THE FIX  --- >>>

            # YOUR ORIGINAL, UNCHANGED FUNCTIONALITY IS BELOW.
            # It will only run on paragraphs that passed the check above (i.e., have no links).
            if find_text in p.text:
                p.text = p.text.replace(find_text, replace_text)

    return doc
